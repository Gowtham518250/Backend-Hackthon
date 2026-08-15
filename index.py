import json
import logging
import os
from pathlib import Path
from typing import Optional

from database import get_rag_documents, get_spreadsheet_rows, rag_file_exists, save_rag_documents, save_spreadsheet_rows

logger = logging.getLogger("deepfind.index")


class RetrievalError(Exception):
    """Raised when every candidate file failed to retrieve (as opposed to
    retrieving successfully and simply finding nothing relevant). Callers
    can catch this specifically to tell 'something is broken' apart from
    'no matching content' instead of both looking like an empty result."""

import pandas as pd
import pymupdf
from PIL import Image
import pytesseract

from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.retrievers import BM25Retriever
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter


SUPPORTED_EXTENSIONS = {
    "pdf",
    "docx",
    "xlsx",
    "xls",
    "csv",
    "jpg",
    "jpeg",
    "png",
    "bmp",
    "tif",
    "tiff",
}

STORE_DIR = Path(
    os.getenv(
        "RAG_DOCUMENT_STORE",
        "RAG_Document_Store",
    )
)

STORE_DIR.mkdir(
    parents=True,
    exist_ok=True,
)

# All document indexing is persisted to PostgreSQL instead of a local FAISS directory.

splitter = RecursiveCharacterTextSplitter(
    chunk_size=int(
        os.getenv("CHUNK_SIZE", "750")
    ),
    chunk_overlap=int(
        os.getenv("CHUNK_OVERLAP", "110")
    ),
)

_embeddings = None


def get_embeddings():
    global _embeddings

    if _embeddings is None:
        model_name = os.getenv(
            "EMBEDDING_MODEL",
            "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
        )

        _embeddings = HuggingFaceEmbeddings(
            model_name=model_name,
            encode_kwargs={
                "normalize_embeddings": True,
            },
        )

    return _embeddings


def file_exists(file_id: str) -> bool:
    return rag_file_exists(file_id)


def detect_language(text: str) -> str:
    if any(
        "\u0B80" <= char <= "\u0BFF"
        for char in text
    ):
        return "ta"

    return "en"


def base_metadata(
    file_id: str,
    original_name: str,
    file_type: str,
):
    return {
        "file_id": file_id,
        "file_name": original_name,
        "file_type": file_type,
    }


def process_pdf(
    file_path: str,
    file_id: str,
    original_name: str,
):
    documents = []

    pdf = pymupdf.open(file_path)

    try:
        for page_number, page in enumerate(
            pdf,
            start=1,
        ):
            text = page.get_text("text").strip()
            ocr = False

            if not text:
                pixmap = page.get_pixmap(
                    dpi=160
                )

                image = Image.frombytes(
                    "RGB",
                    [
                        pixmap.width,
                        pixmap.height,
                    ],
                    pixmap.samples,
                )

                text = pytesseract.image_to_string(
                    image
                ).strip()

                ocr = True

            if not text:
                continue

            metadata = base_metadata(
                file_id,
                original_name,
                "pdf",
            )

            metadata.update({
                "page": page_number,
                "language": detect_language(text),
                "ocr": ocr,
                "chunk_type": "page",
            })

            documents.append(
                Document(
                    page_content=text,
                    metadata=metadata,
                )
            )

    finally:
        pdf.close()

    return documents


def process_docx(
    file_path: str,
    file_id: str,
    original_name: str,
):
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    documents = []

    for index, paragraph in enumerate(
        doc.paragraphs,
        start=1,
    ):
        text = paragraph.text.strip()

        if not text:
            continue

        metadata = base_metadata(
            file_id,
            original_name,
            "docx",
        )

        metadata.update({
            "paragraph": index,
            "section": (
                paragraph.style.name
                if paragraph.style
                else None
            ),
            "language": detect_language(text),
            "chunk_type": "paragraph",
        })

        documents.append(
            Document(
                page_content=text,
                metadata=metadata,
            )
        )

    # Also index DOCX tables.
    for table_index, table in enumerate(
        doc.tables,
        start=1,
    ):
        for row_index, row in enumerate(
            table.rows,
            start=1,
        ):
            values = [
                cell.text.strip()
                for cell in row.cells
            ]

            text = " | ".join(values)

            if not text.strip():
                continue

            metadata = base_metadata(
                file_id,
                original_name,
                "docx",
            )

            metadata.update({
                "table": table_index,
                "row": row_index,
                "language": detect_language(text),
                "chunk_type": "table_row",
            })

            documents.append(
                Document(
                    page_content=text,
                    metadata=metadata,
                )
            )

    return documents


def process_excel(
    file_path: str,
    file_id: str,
    original_name: str,
):
    documents = []

    workbook = pd.ExcelFile(file_path)
    try:
        for sheet_name in workbook.sheet_names:
            dataframe = pd.read_excel(
                file_path,
                sheet_name=sheet_name,
            )

        headers = [
            str(column)
            for column in dataframe.columns
        ]

        for index, row in dataframe.iterrows():
            values = {}

            for column in dataframe.columns:
                value = row[column]

                if pd.isna(value):
                    continue

                if hasattr(
                    value,
                    "isoformat",
                ):
                    value = value.isoformat()

                values[str(column)] = value

            if not values:
                continue

            save_spreadsheet_rows(
                file_id=file_id,
                file_name=original_name,
                sheet_name=sheet_name,
                headers=headers,
                row_data=values,
            )

            text = " | ".join(
                f"{key}: {value}"
                for key, value in values.items()
            )

            metadata = base_metadata(
                file_id,
                original_name,
                "xlsx",
            )

            metadata.update({
                "sheet": sheet_name,
                "row": int(index) + 2,
                "headers": headers,
                "row_data": values,
                "language": detect_language(text),
                "chunk_type": "structured_row",
            })

            documents.append(
                Document(
                    page_content=text,
                    metadata=metadata,
                )
            )
    finally:
        workbook.close()

    return documents


def process_csv(
    file_path: str,
    file_id: str,
    original_name: str,
):
    dataframe = pd.read_csv(file_path)

    documents = []

    headers = [
        str(column)
        for column in dataframe.columns
    ]

    for index, row in dataframe.iterrows():
        values = {}

        for column in dataframe.columns:
            value = row[column]

            if pd.isna(value):
                continue

            if hasattr(
                value,
                "isoformat",
            ):
                value = value.isoformat()

            values[str(column)] = value

        if not values:
            continue

        save_spreadsheet_rows(
            file_id=file_id,
            file_name=original_name,
            sheet_name="csv",
            headers=headers,
            row_data=values,
        )

        text = " | ".join(
            f"{key}: {value}"
            for key, value in values.items()
        )

        metadata = base_metadata(
            file_id,
            original_name,
            "csv",
        )

        metadata.update({
            "row": int(index) + 2,
            "headers": headers,
            "row_data": values,
            "language": detect_language(text),
            "chunk_type": "structured_row",
        })

        documents.append(
            Document(
                page_content=text,
                metadata=metadata,
            )
        )

    return documents


def process_image(
    file_path: str,
    file_id: str,
    original_name: str,
):
    image = Image.open(file_path)

    text = pytesseract.image_to_string(
        image
    ).strip()

    if not text:
        return []

    metadata = base_metadata(
        file_id,
        original_name,
        "image",
    )

    metadata.update({
        "ocr": True,
        "language": detect_language(text),
        "chunk_type": "ocr",
    })

    return [
        Document(
            page_content=text,
            metadata=metadata,
        )
    ]


def extract_documents(
    file_path: str,
    file_id: str,
    original_name: str,
):
    extension = (
        Path(file_path)
        .suffix
        .lower()
        .lstrip(".")
    )

    if extension == "pdf":
        return process_pdf(
            file_path,
            file_id,
            original_name,
        )

    if extension == "docx":
        return process_docx(
            file_path,
            file_id,
            original_name,
        )

    if extension in {"xlsx", "xls"}:
        return process_excel(
            file_path,
            file_id,
            original_name,
        )

    if extension == "csv":
        return process_csv(
            file_path,
            file_id,
            original_name,
        )

    if extension in {
        "jpg",
        "jpeg",
        "png",
        "bmp",
        "tif",
        "tiff",
    }:
        return process_image(
            file_path,
            file_id,
            original_name,
        )

    raise ValueError(
        f"Unsupported file type: .{extension}"
    )


def register_file(
    file_path: str,
    file_id: str,
    original_name: Optional[str] = None,
):
    original_name = (
        original_name
        or Path(file_path).name
    )

    documents = extract_documents(
        file_path=file_path,
        file_id=file_id,
        original_name=original_name,
    )

    if not documents:
        raise ValueError(
            "No searchable content was extracted."
        )

    chunks = splitter.split_documents(
        documents
    )

    for index, chunk in enumerate(chunks):
        chunk.metadata["chunk_id"] = (
            f"{file_id}-{index}"
        )
        chunk.metadata["chunk_index"] = index

    serialized = [
        {
            "page_content": chunk.page_content,
            "metadata": chunk.metadata,
        }
        for chunk in chunks
    ]

    save_rag_documents(
        file_id=file_id,
        file_name=original_name,
        file_type=(Path(file_path).suffix.lower().lstrip(".")),
        documents=serialized,
    )

    return {
        "file_id": file_id,
        "file_name": original_name,
        "chunks": len(chunks),
        "file_type": (
            Path(file_path)
            .suffix
            .lower()
            .lstrip(".")
        ),
    }


def load_documents(
    file_id: str,
):
    raw = get_rag_documents(file_id=file_id)

    if not raw:
        raise FileNotFoundError(
            f"Indexed documents missing: {file_id}"
        )

    return [
        Document(
            page_content=item["page_content"],
            metadata=item.get(
                "metadata",
                {},
            ),
        )
        for item in raw
    ]


def access_file(file_id: str):
    if not file_exists(file_id):
        raise FileNotFoundError(
            f"Indexed file not found: {file_id}"
        )

    documents = load_documents(file_id)

    bm25 = BM25Retriever.from_documents(documents)
    bm25.k = int(os.getenv("BM25_K", "7"))

    faiss_store = FAISS.from_documents(
        documents=documents,
        embedding=get_embeddings(),
    )

    return HybridRetriever(
        bm25=bm25,
        faiss_store=faiss_store,
    )


class HybridRetriever:
    def __init__(
        self,
        bm25,
        faiss_store,
    ):
        self.bm25 = bm25
        self.faiss_store = faiss_store

    @staticmethod
    def document_key(document):
        metadata = document.metadata or {}

        return metadata.get(
            "chunk_id",
            (
                f"{metadata.get('file_name', '')}:"
                f"{document.page_content}"
            ),
        )

    def invoke(
        self,
        query: str,
        top_k: int = 8,
    ):
        keyword_docs = self.bm25.invoke(
            query
        )

        semantic_docs = (
            self.faiss_store.similarity_search(
                query,
                k=max(top_k, 8),
            )
        )

        scores = {}
        documents = {}

        # Weighted reciprocal-rank fusion.
        for rank, document in enumerate(
            keyword_docs,
            start=1,
        ):
            key = self.document_key(
                document
            )

            documents[key] = document

            scores[key] = (
                scores.get(key, 0.0)
                + 0.4 / (60 + rank)
            )

        for rank, document in enumerate(
            semantic_docs,
            start=1,
        ):
            key = self.document_key(
                document
            )

            documents[key] = document

            scores[key] = (
                scores.get(key, 0.0)
                + 0.6 / (60 + rank)
            )

        ordered = sorted(
            documents,
            key=lambda key: scores[key],
            reverse=True,
        )

        return [
            documents[key]
            for key in ordered[:top_k]
        ]


def search_all_files(
    query: str,
    top_k: int = 8,
    file_id: Optional[str] = None,
):
    if file_id:
        return access_file(
            file_id
        ).invoke(
            query,
            top_k=top_k,
        )

    candidates = []

    rows = get_rag_documents()
    file_ids = []
    for row in rows:
        # NOTE: this used to reuse the outer `file_id` parameter name,
        # which happened to be harmless only because we already know it's
        # None on this branch — renamed to avoid confusing future edits.
        row_file_id = (row.get("metadata") or {}).get("file_id")
        if row_file_id and row_file_id not in file_ids:
            file_ids.append(row_file_id)

    errors = []
    for current_id in file_ids:
        try:
            retriever = access_file(current_id)
            candidates.extend(retriever.invoke(query, top_k=top_k))
        except Exception as exc:
            logger.exception("Retrieval failed for file_id=%s", current_id)
            errors.append((current_id, exc))

    if file_ids and not candidates and len(errors) == len(file_ids):
        # Every single file failed to retrieve — this is a broken pipeline
        # (e.g. embedding model couldn't load), not "no matching content".
        # Surface it instead of silently returning an empty list.
        raise RetrievalError(
            f"All {len(errors)} file(s) failed during retrieval. "
            f"First error for file_id={errors[0][0]}: {errors[0][1]}"
        )

    # Global merge across all indexed files.
    scores = {}
    documents = {}

    for rank, document in enumerate(
        candidates,
        start=1,
    ):
        key = (
            document.metadata or {}
        ).get(
            "chunk_id",
            f"unknown-{rank}",
        )

        documents[key] = document

        scores[key] = (
            scores.get(key, 0.0)
            + 1.0 / (60 + rank)
        )

    ordered = sorted(
        documents,
        key=lambda key: scores[key],
        reverse=True,
    )

    return [
        documents[key]
        for key in ordered[:top_k]
    ]